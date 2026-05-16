"""PDF/DOCX экспорт КП.

TipTap JSON → HTML/DOCX walker. Поддерживает кастомные ноды (MergeTag,
PaymentSchedule) с резолвом данных из get_merge_data_internal.

PDF — через WeasyPrint (уже стоит в ERPNext).
DOCX — через python-docx (программный walker).
"""
from __future__ import annotations

import json
from io import BytesIO

import frappe

from olimp_construction.api.proposals import get_merge_data_internal


# ─────────────────────────── Merge tag resolver ─────────────────────────────


def _resolve_merge_tag(path: str, data: dict) -> str:
    """customer.name → data['customer']['name']."""
    if not path:
        return ""
    parts = path.split(".")
    val: object = data
    try:
        for p in parts:
            val = val[p]  # type: ignore[index]
    except (KeyError, TypeError):
        return f"{{{{{path}}}}}"
    if val is None:
        return ""
    if isinstance(val, (int, float)):
        # Числа форматируем с разделителями тысяч
        return f"{val:,.0f}".replace(",", " ")
    return str(val)


# ─────────────────────────── TipTap JSON → HTML ─────────────────────────────


def _escape(s: str) -> str:
    return (str(s).replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


def _render_marks(text: str, marks: list[dict] | None) -> str:
    """Применяет marks (bold/italic/strike) к тексту."""
    out = _escape(text)
    for m in marks or []:
        t = m.get("type")
        if t == "bold":
            out = f"<strong>{out}</strong>"
        elif t == "italic":
            out = f"<em>{out}</em>"
        elif t == "strike":
            out = f"<s>{out}</s>"
        elif t == "code":
            out = f"<code>{out}</code>"
        elif t == "link":
            href = m.get("attrs", {}).get("href", "#")
            out = f'<a href="{_escape(href)}">{out}</a>'
    return out


def _render_node(node: dict, merge: dict) -> str:
    """Рекурсивно рендерит TipTap node в HTML."""
    if not isinstance(node, dict):
        return ""
    t = node.get("type")
    attrs = node.get("attrs", {}) or {}
    content = node.get("content", []) or []

    def kids() -> str:
        return "".join(_render_node(c, merge) for c in content)

    if t == "doc":
        return kids()
    if t == "paragraph":
        inner = kids()
        return f"<p>{inner or '&nbsp;'}</p>"
    if t == "text":
        return _render_marks(node.get("text", ""), node.get("marks"))
    if t == "heading":
        level = attrs.get("level", 1)
        return f"<h{level}>{kids()}</h{level}>"
    if t == "bulletList":
        return f"<ul>{kids()}</ul>"
    if t == "orderedList":
        return f"<ol>{kids()}</ol>"
    if t == "listItem":
        return f"<li>{kids()}</li>"
    if t == "blockquote":
        return f"<blockquote>{kids()}</blockquote>"
    if t == "horizontalRule":
        return "<hr/>"
    if t == "hardBreak":
        return "<br/>"
    if t == "codeBlock":
        return f"<pre><code>{kids()}</code></pre>"

    if t == "table":
        return f'<table class="tt-table">{kids()}</table>'
    if t == "tableRow":
        return f"<tr>{kids()}</tr>"
    if t == "tableHeader":
        return f"<th>{kids()}</th>"
    if t == "tableCell":
        return f"<td>{kids()}</td>"

    if t == "image":
        src = attrs.get("src", "")
        alt = attrs.get("alt", "")
        return f'<img src="{_escape(src)}" alt="{_escape(alt)}" style="max-width:100%"/>'

    # ── Кастомные ноды ────────────────────────────────────────────────────
    if t == "mergeTag":
        return f'<span class="merge-tag">{_escape(_resolve_merge_tag(attrs.get("path", ""), merge))}</span>'

    if t == "paymentSchedule":
        return _render_payment_schedule(attrs, merge)

    if t == "spreadsheet":
        return _render_spreadsheet_snapshot(attrs)

    # Fallback — рекурсия по детям
    return kids()


def _render_payment_schedule(attrs: dict, merge: dict) -> str:
    rows = attrs.get("rows", []) or []
    base_amount = float(attrs.get("total_amount") or 0)
    if not base_amount:
        try:
            base_amount = float(merge.get("proposal", {}).get("total_amount") or 0)
        except (TypeError, ValueError):
            base_amount = 0

    parts = ['<table class="payment-schedule"><thead><tr>',
             '<th>Этап</th><th>%</th><th>Срок</th><th>Сумма ₽</th></tr></thead><tbody>']
    total_pct = 0.0
    for r in rows:
        stage = _escape(str(r.get("stage", "")))
        pct = float(r.get("percent", 0) or 0)
        days = int(r.get("days_after", 0) or 0)
        amount = base_amount * pct / 100 if base_amount else 0
        total_pct += pct
        amount_str = f"{amount:,.0f}".replace(",", " ") if amount else "—"
        days_str = "при подписании" if days == 0 else f"+{days} дн."
        parts.append(f"<tr><td>{stage}</td><td>{pct:.1f}%</td><td>{days_str}</td>"
                     f"<td style='text-align:right'>{amount_str}</td></tr>")
    total_str = f"{base_amount:,.0f}".replace(",", " ") if base_amount else "—"
    parts.append(f"<tr class='tt-total'><td><b>ИТОГО</b></td><td><b>{total_pct:.1f}%</b></td><td></td>"
                 f"<td style='text-align:right'><b>{total_str}</b></td></tr>")
    parts.append("</tbody></table>")
    return "".join(parts)


def _render_spreadsheet_snapshot(attrs: dict) -> str:
    """Univer snapshot → простая HTML-таблица (для PDF render)."""
    snapshot = attrs.get("snapshot") or {}
    if not snapshot:
        return '<div class="spreadsheet-placeholder">[Spreadsheet block]</div>'

    # Структура Univer snapshot: sheets[sheetId].cellData[rowIdx][colIdx] = {v: value}
    try:
        sheets = snapshot.get("sheets", {}) or {}
        first_sheet = next(iter(sheets.values()), {}) if sheets else {}
        cell_data = first_sheet.get("cellData", {}) or {}
        if not cell_data:
            return '<div class="spreadsheet-placeholder">[Пустая таблица]</div>'

        max_row = max(int(r) for r in cell_data.keys())
        max_col = 0
        for r in cell_data.values():
            if r:
                max_col = max(max_col, max(int(c) for c in r.keys()))

        parts = ['<table class="tt-table">']
        for ri in range(max_row + 1):
            parts.append("<tr>")
            row = cell_data.get(str(ri), {}) or {}
            for ci in range(max_col + 1):
                cell = row.get(str(ci), {}) or {}
                v = cell.get("v", "")
                parts.append(f"<td>{_escape(str(v))}</td>")
            parts.append("</tr>")
        parts.append("</table>")
        return "".join(parts)
    except Exception:
        return '<div class="spreadsheet-placeholder">[Spreadsheet]</div>'


def tiptap_to_html(content: dict, merge: dict) -> str:
    """Главный entry point для TipTap JSON → HTML."""
    if not content or not isinstance(content, dict):
        return ""
    return _render_node(content, merge)


# ─────────────────────────── Полноценный HTML wrapper ───────────────────────


_BASE_CSS = """
@page { size: A4; margin: 18mm 16mm; }
* { box-sizing: border-box; }
body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #1d1d1f;
       font-size: 11.5pt; line-height: 1.5; margin: 0; }
.header { border-bottom: 2px solid #1d1d1f; padding-bottom: 12px;
          margin-bottom: 24px; }
.header h1 { font-size: 22pt; font-weight: 600; margin: 0 0 4px; }
.header .meta { font-size: 10pt; color: #666; }
.kp-summary { display: flex; justify-content: space-between; margin: 12px 0 20px;
              padding: 12px 16px; background: #f5f5f7; border-radius: 8px; }
.kp-summary .label { font-size: 9pt; color: #666; text-transform: uppercase;
                     letter-spacing: 0.5pt; }
.kp-summary .value { font-size: 14pt; font-weight: 600; }
h1, h2, h3 { color: #1d1d1f; margin-top: 18pt; margin-bottom: 8pt;
             font-weight: 600; }
h1 { font-size: 18pt; }
h2 { font-size: 14pt; }
h3 { font-size: 12pt; }
p { margin: 6pt 0; }
ul, ol { margin: 6pt 0; padding-left: 22pt; }
blockquote { margin: 10pt 0; padding: 8pt 14pt; border-left: 3pt solid #d2d2d7;
             color: #555; font-style: italic; }
.merge-tag { background: rgba(74,222,128,0.15); padding: 1pt 4pt;
             border-radius: 3pt; color: #16a34a; font-weight: 500; }
table.tt-table { border-collapse: collapse; width: 100%; margin: 10pt 0;
                 font-size: 10pt; }
table.tt-table td, table.tt-table th { border: 0.5pt solid #d2d2d7;
                                        padding: 5pt 8pt; vertical-align: top; }
table.tt-table th { background: #f5f5f7; font-weight: 600; }
table.payment-schedule { border-collapse: collapse; width: 100%;
                         margin: 14pt 0; font-size: 10pt; }
table.payment-schedule th { background: #1d1d1f; color: white;
                            padding: 6pt 10pt; text-align: left; }
table.payment-schedule td { border: 0.5pt solid #d2d2d7;
                            padding: 5pt 10pt; }
table.payment-schedule tr.tt-total td { background: #f5f5f7; font-weight: 600;
                                         border-top: 1pt solid #1d1d1f; }
.footer { margin-top: 30pt; padding-top: 14pt; border-top: 1pt solid #d2d2d7;
          font-size: 9pt; color: #888; text-align: center; }
.sign-block { margin-top: 30pt; padding: 14pt; border: 1pt solid #d2d2d7;
              border-radius: 6pt; }
.sign-block img { max-height: 80pt; }
"""


def wrap_html(body_html: str, doc) -> str:
    """Оборачивает тело КП в полноценный HTML документ для PDF/print."""
    merge = get_merge_data_internal(doc)
    customer = merge["customer"]["name"] or "—"
    project = merge["project"]["title"] or ""
    valid_until = str(doc.valid_until) if doc.valid_until else ""
    total = float(doc.total_amount or 0)

    sign_html = ""
    if doc.signed_by_name and doc.signature_data_url:
        sign_html = f"""
        <div class="sign-block">
          <div style="font-weight:600;margin-bottom:6pt;">✓ Согласовано клиентом</div>
          <div style="font-size:10pt;color:#666;margin-bottom:8pt;">
            {_escape(doc.signed_by_name)} — {str(doc.signed_at).split('.')[0] if doc.signed_at else ''}
          </div>
          <img src="{_escape(doc.signature_data_url)}" alt="Подпись"/>
        </div>
        """

    return f"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8"/>
<title>{_escape(doc.title)} — {_escape(doc.name)}</title>
<style>{_BASE_CSS}</style>
</head><body>
<div class="header">
  <h1>{_escape(doc.title)}</h1>
  <div class="meta">{_escape(doc.name)} · {merge['proposal']['today']}</div>
</div>
<div class="kp-summary">
  <div>
    <div class="label">Заказчик</div>
    <div class="value" style="font-size:11pt;">{_escape(customer)}</div>
    {f'<div style="font-size:9pt;color:#888;margin-top:2pt;">{_escape(project)}</div>' if project else ''}
  </div>
  <div style="text-align:right;">
    <div class="label">Сумма КП</div>
    <div class="value">{total:,.0f} ₽</div>
    {f'<div style="font-size:9pt;color:#888;margin-top:2pt;">Действует до: {valid_until}</div>' if valid_until else ''}
  </div>
</div>
{body_html}
{sign_html}
<div class="footer">ООО «Олимп» · Промышленное строительство · Екатеринбург</div>
</body></html>""".replace(",", " ")


# ─────────────────────────── PDF export ─────────────────────────────────────


@frappe.whitelist(allow_guest=False, methods=["GET"])
def export_pdf(name: str) -> None:
    """Сгенерировать PDF и вернуть как download."""
    frappe.has_permission("Construction Proposal", "read", doc=name, throw=True)
    doc = frappe.get_doc("Construction Proposal", name)

    try:
        content = json.loads(doc.content_json or "{}")
    except (json.JSONDecodeError, TypeError):
        content = {}

    merge = get_merge_data_internal(doc)
    body_html = tiptap_to_html(content, merge)
    full_html = wrap_html(body_html, doc)

    # WeasyPrint
    try:
        from weasyprint import HTML
        pdf_bytes = HTML(string=full_html).write_pdf()
    except Exception as e:
        frappe.throw(f"PDF generator не доступен: {e}")

    safe_name = doc.name.replace("/", "-")
    frappe.local.response.filename = f"{safe_name}.pdf"
    frappe.local.response.filecontent = pdf_bytes
    frappe.local.response.type = "download"


# ─────────────────────────── DOCX export ────────────────────────────────────


def _docx_walk(doc_obj, node: dict, merge: dict, doc_props) -> None:
    """Рекурсивный walker по TipTap node → docx document."""
    if not isinstance(node, dict):
        return
    t = node.get("type")
    attrs = node.get("attrs", {}) or {}
    content = node.get("content", []) or []

    if t == "doc":
        for c in content:
            _docx_walk(doc_obj, c, merge, doc_props)
        return

    if t == "paragraph":
        p = doc_obj.add_paragraph()
        for c in content:
            _docx_inline(p, c, merge)
        return

    if t == "heading":
        level = attrs.get("level", 1)
        text_parts = []
        for c in content:
            if c.get("type") == "text":
                text_parts.append(c.get("text", ""))
            elif c.get("type") == "mergeTag":
                text_parts.append(_resolve_merge_tag(c.get("attrs", {}).get("path", ""), merge))
        doc_obj.add_heading("".join(text_parts), level=min(level, 4))
        return

    if t == "bulletList":
        for li in content:
            _docx_list_item(doc_obj, li, merge, "List Bullet")
        return
    if t == "orderedList":
        for li in content:
            _docx_list_item(doc_obj, li, merge, "List Number")
        return

    if t == "blockquote":
        for c in content:
            _docx_walk(doc_obj, c, merge, doc_props)
        return

    if t == "table":
        rows = [c for c in content if c.get("type") == "tableRow"]
        if not rows:
            return
        max_cols = max(
            len([cell for cell in r.get("content", []) if cell.get("type") in ("tableCell", "tableHeader")])
            for r in rows
        )
        if max_cols == 0:
            return
        table = doc_obj.add_table(rows=len(rows), cols=max_cols)
        table.style = "Light Grid"
        for ri, r in enumerate(rows):
            cells = [c for c in r.get("content", []) if c.get("type") in ("tableCell", "tableHeader")]
            for ci, cell_node in enumerate(cells[:max_cols]):
                cell = table.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                for inner in cell_node.get("content", []):
                    if inner.get("type") == "paragraph":
                        for span in inner.get("content", []):
                            _docx_inline(p, span, merge)
        return

    if t == "paymentSchedule":
        _docx_payment_schedule(doc_obj, attrs, merge)
        return

    if t == "spreadsheet":
        _docx_spreadsheet(doc_obj, attrs)
        return

    if t == "horizontalRule":
        doc_obj.add_paragraph("─" * 50)
        return

    # Fallback — рекурсия по детям
    for c in content:
        _docx_walk(doc_obj, c, merge, doc_props)


def _docx_inline(paragraph, node: dict, merge: dict) -> None:
    """Inline node → run в paragraph."""
    if not isinstance(node, dict):
        return
    t = node.get("type")
    if t == "text":
        text = node.get("text", "")
        marks = {m.get("type") for m in node.get("marks", []) or []}
        run = paragraph.add_run(text)
        if "bold" in marks: run.bold = True
        if "italic" in marks: run.italic = True
        if "strike" in marks:
            run.font.strike = True
        if "code" in marks:
            run.font.name = "Courier New"
        return
    if t == "mergeTag":
        resolved = _resolve_merge_tag(node.get("attrs", {}).get("path", ""), merge)
        paragraph.add_run(resolved).bold = True
        return
    if t == "hardBreak":
        paragraph.add_run().add_break()
        return


def _docx_list_item(doc_obj, li: dict, merge: dict, style: str) -> None:
    for child in li.get("content", []):
        if child.get("type") == "paragraph":
            p = doc_obj.add_paragraph(style=style)
            for span in child.get("content", []):
                _docx_inline(p, span, merge)


def _docx_payment_schedule(doc_obj, attrs: dict, merge: dict) -> None:
    rows = attrs.get("rows", []) or []
    base_amount = float(attrs.get("total_amount") or 0)
    if not base_amount:
        try:
            base_amount = float(merge.get("proposal", {}).get("total_amount") or 0)
        except (TypeError, ValueError):
            base_amount = 0

    doc_obj.add_heading("График оплаты", level=3)
    table = doc_obj.add_table(rows=len(rows) + 2, cols=4)
    table.style = "Light Grid"

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(["Этап", "%", "Срок", "Сумма ₽"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    total_pct = 0.0
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        pct = float(row.get("percent", 0) or 0)
        days = int(row.get("days_after", 0) or 0)
        amount = base_amount * pct / 100 if base_amount else 0
        total_pct += pct
        cells[0].text = str(row.get("stage", ""))
        cells[1].text = f"{pct:.1f}%"
        cells[2].text = "при подписании" if days == 0 else f"+{days} дн."
        cells[3].text = f"{amount:,.0f}".replace(",", " ") if amount else "—"

    # Total row
    total_cells = table.rows[-1].cells
    total_cells[0].text = "ИТОГО"
    total_cells[1].text = f"{total_pct:.1f}%"
    total_cells[3].text = f"{base_amount:,.0f}".replace(",", " ") if base_amount else "—"
    for c in total_cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True


def _docx_spreadsheet(doc_obj, attrs: dict) -> None:
    """Univer snapshot → docx table (плоско)."""
    snapshot = attrs.get("snapshot") or {}
    if not snapshot:
        doc_obj.add_paragraph("[Spreadsheet block]")
        return
    try:
        sheets = snapshot.get("sheets", {}) or {}
        first_sheet = next(iter(sheets.values()), {}) if sheets else {}
        cell_data = first_sheet.get("cellData", {}) or {}
        if not cell_data:
            doc_obj.add_paragraph("[Пустая таблица]")
            return
        max_row = max(int(r) for r in cell_data.keys())
        max_col = 0
        for r in cell_data.values():
            if r:
                max_col = max(max_col, max(int(c) for c in r.keys()))

        table = doc_obj.add_table(rows=max_row + 1, cols=max_col + 1)
        table.style = "Light Grid"
        for ri in range(max_row + 1):
            row = cell_data.get(str(ri), {}) or {}
            for ci in range(max_col + 1):
                cell = row.get(str(ci), {}) or {}
                table.cell(ri, ci).text = str(cell.get("v", ""))
    except Exception:
        doc_obj.add_paragraph("[Spreadsheet]")


@frappe.whitelist(allow_guest=False, methods=["GET"])
def export_docx(name: str) -> None:
    """Сгенерировать DOCX и вернуть как download."""
    frappe.has_permission("Construction Proposal", "read", doc=name, throw=True)
    doc = frappe.get_doc("Construction Proposal", name)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        frappe.throw("python-docx не установлен. pip install python-docx")

    try:
        content = json.loads(doc.content_json or "{}")
    except (json.JSONDecodeError, TypeError):
        content = {}

    merge = get_merge_data_internal(doc)

    # Build docx
    docx_doc = Document()

    # Default font
    for style_name in ("Normal", "List Bullet", "List Number"):
        try:
            st = docx_doc.styles[style_name]
            st.font.name = "Calibri"
            st.font.size = Pt(11)
        except KeyError:
            pass

    # Header
    h = docx_doc.add_heading(doc.title or "Коммерческое предложение", level=0)

    info_p = docx_doc.add_paragraph()
    info_p.add_run(f"{doc.name} · {merge['proposal']['today']}").font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Customer/total info table
    info_table = docx_doc.add_table(rows=2, cols=2)
    info_table.style = "Light Grid"
    info_table.rows[0].cells[0].text = "Заказчик"
    info_table.rows[0].cells[1].text = merge["customer"]["name"] or "—"
    info_table.rows[1].cells[0].text = "Сумма КП"
    info_table.rows[1].cells[1].text = f"{float(doc.total_amount or 0):,.0f} ₽".replace(",", " ")
    if doc.valid_until:
        row = info_table.add_row().cells
        row[0].text = "Действует до"
        row[1].text = str(doc.valid_until)
    if merge["project"]["title"]:
        row = info_table.add_row().cells
        row[0].text = "Проект"
        row[1].text = merge["project"]["title"]

    docx_doc.add_paragraph()  # spacer

    # Walk through content
    _docx_walk(docx_doc, content, merge, None)

    # Signature info
    if doc.signed_by_name:
        docx_doc.add_paragraph()
        sign_p = docx_doc.add_paragraph()
        sign_run = sign_p.add_run(f"✓ Согласовано клиентом: {doc.signed_by_name}")
        sign_run.bold = True
        if doc.signed_at:
            sign_p.add_run(f" — {str(doc.signed_at).split('.')[0]}")

    # Footer
    footer_p = docx_doc.add_paragraph()
    footer_p.alignment = 1  # CENTER
    footer_run = footer_p.add_run("ООО «Олимп» · Промышленное строительство · Екатеринбург")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    docx_doc.save(buf)
    buf.seek(0)

    safe_name = doc.name.replace("/", "-")
    frappe.local.response.filename = f"{safe_name}.docx"
    frappe.local.response.filecontent = buf.getvalue()
    frappe.local.response.type = "download"


# ─────────────────────────── Preview HTML ───────────────────────────────────


@frappe.whitelist()
def preview_html(name: str) -> dict:
    """HTML preview КП (для модалки превью)."""
    frappe.has_permission("Construction Proposal", "read", doc=name, throw=True)
    doc = frappe.get_doc("Construction Proposal", name)
    try:
        content = json.loads(doc.content_json or "{}")
    except (json.JSONDecodeError, TypeError):
        content = {}
    merge = get_merge_data_internal(doc)
    body = tiptap_to_html(content, merge)
    full = wrap_html(body, doc)
    return {"html": full}
